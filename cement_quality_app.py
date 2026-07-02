"""
Cement Clinker Quality Analyzer
================================
A comprehensive Streamlit app for cement plant quality control.
Handles multi-sheet Excel files, image OCR, trend analysis, and report generation.
Works fully offline.
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import cv2
from PIL import Image
import pytesseract
import io
import re
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import tempfile

# =============================================================================
# PAGE CONFIGURATION
# =============================================================================
st.set_page_config(
    page_title="Cement Quality Analyzer",
    page_icon="🏭",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================================
# CUSTOM CSS FOR BETTER UI
# =============================================================================
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 800;
        color: #1f4e79;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.1rem;
        color: #555;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 1rem;
        border-left: 5px solid #1f4e79;
    }
    .insight-box {
        background-color: #e8f4f8;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
        border-left: 4px solid #17a2b8;
    }
    .alert-box {
        background-color: #fff3cd;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
        border-left: 4px solid #ffc107;
    }
    .danger-box {
        background-color: #f8d7da;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
        border-left: 4px solid #dc3545;
    }
    .success-box {
        background-color: #d4edda;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
        border-left: 4px solid #28a745;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #f0f2f6;
        border-radius: 8px 8px 0 0;
        gap: 1px;
        padding-top: 10px;
        padding-bottom: 10px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #1f4e79;
        color: white;
    }
    div[data-testid="stFileUploader"] {
        border: 2px dashed #1f4e79;
        border-radius: 10px;
        padding: 20px;
    }
    .copy-btn {
        background-color: #4CAF50;
        color: white;
        border: none;
        padding: 8px 16px;
        border-radius: 5px;
        cursor: pointer;
        font-weight: 600;
    }
    .copy-btn:hover {
        background-color: #45a049;
    }
</style>
""", unsafe_allow_html=True)

# =============================================================================
# MATERIAL DETECTION & SPECIFICATIONS
# =============================================================================

MATERIAL_KEYWORDS = {
    'Kiln Feed': {
        'keywords': ['kiln feed', 'kilnfeed', 'raw meal', 'feed', 'lsf', 'sm', 'am', 
                     'moisture', 'loi', 'lime saturation', 'silica modulus', 'alumina modulus',
                     'kh', 'n', 'p', 'homogeneity'],
        'params': {
            'LSF': {'min': 90, 'max': 110, 'unit': '%'},
            'SM': {'min': 2.0, 'max': 3.0, 'unit': ''},
            'AM': {'min': 1.0, 'max': 1.8, 'unit': ''},
            'Moisture': {'min': 0, 'max': 1.0, 'unit': '%'},
            'LOI': {'min': 30, 'max': 40, 'unit': '%'}
        }
    },
    'Clinker': {
        'keywords': ['clinker', 'c3s', 'c2s', 'c3a', 'c4af', 'free lime', 'f.cao', 
                     'fcao', 'free cao', 'liter weight', 'litrature', 'literweight',
                     'alite', 'belite', 'aluminate', 'ferrite'],
        'params': {
            'C3S': {'min': 50, 'max': 70, 'unit': '%'},
            'C2S': {'min': 10, 'max': 30, 'unit': '%'},
            'C3A': {'min': 5, 'max': 12, 'unit': '%'},
            'C4AF': {'min': 5, 'max': 12, 'unit': '%'},
            'Free Lime': {'min': 0, 'max': 1.5, 'unit': '%'},
            'Liter Weight': {'min': 1200, 'max': 1400, 'unit': 'g/L'}
        }
    },
    'Cement': {
        'keywords': ['cement', 'blaine', 'so3', 'sulfur trioxide', 'setting time', 
                     'compressive', 'strength', 'fineness', 'retention', '45 micron',
                     'initial setting', 'final setting', '3 day', '7 day', '28 day',
                     'soundness', 'autoclave', 'loi'],
        'params': {
            'Blaine': {'min': 300, 'max': 450, 'unit': 'm²/kg'},
            'SO3': {'min': 0, 'max': 3.5, 'unit': '%'},
            '3-Day Strength': {'min': 20, 'max': 35, 'unit': 'MPa'},
            '7-Day Strength': {'min': 30, 'max': 45, 'unit': 'MPa'},
            '28-Day Strength': {'min': 40, 'max': 60, 'unit': 'MPa'},
            'Initial Setting': {'min': 60, 'max': 180, 'unit': 'min'},
            'Final Setting': {'min': 180, 'max': 360, 'unit': 'min'}
        }
    }
}

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def detect_material_type(df, sheet_name=""):
    """Detect material type based on column headers and sheet name."""
    cols = ' '.join(df.columns.astype(str)).lower()
    name = sheet_name.lower()
    combined = cols + " " + name

    scores = {}
    for material, info in MATERIAL_KEYWORDS.items():
        score = 0
        for kw in info['keywords']:
            if kw in combined:
                score += 1
        scores[material] = score

    if max(scores.values()) == 0:
        # Try row-based detection if headers didn't match
        sample_text = df.head(10).to_string().lower()
        for material, info in MATERIAL_KEYWORDS.items():
            score = 0
            for kw in info['keywords']:
                if kw in sample_text:
                    score += 1
            scores[material] = score

    best_match = max(scores, key=scores.get)
    if scores[best_match] == 0:
        return "Unknown"
    return best_match


def clean_dataframe(df):
    """Clean and prepare dataframe for analysis."""
    # Drop completely empty rows and columns
    df = df.dropna(how='all').dropna(axis=1, how='all')

    # Try to detect header row if first row looks like data
    if len(df) > 1:
        first_row = df.iloc[0].astype(str)
        if any(keyword in ' '.join(first_row).lower() for keyword in 
               ['date', 'sample', 'lsf', 'c3s', 'blaine', 'strength', 'so3']):
            df.columns = first_row
            df = df.iloc[1:].reset_index(drop=True)

    # Convert numeric columns
    for col in df.columns:
        df[col] = pd.to_numeric(df[col], errors='ignore')

    return df.reset_index(drop=True)


def analyze_sheet(df, material_type):
    """Perform statistical analysis on a sheet."""
    if df.empty:
        return None

    # Get numeric columns only for analysis
    numeric_df = df.select_dtypes(include=[np.number])

    if numeric_df.empty:
        return None

    stats = {
        'material': material_type,
        'row_count': len(df),
        'columns': list(df.columns),
        'numeric_summary': numeric_df.describe().to_dict(),
        'out_of_spec': [],
        'trends': {}
    }

    # Check specifications
    specs = MATERIAL_KEYWORDS.get(material_type, {}).get('params', {})
    for param, limits in specs.items():
        # Find matching column (fuzzy)
        for col in numeric_df.columns:
            if param.lower() in str(col).lower() or str(col).lower() in param.lower():
                values = numeric_df[col].dropna()
                if len(values) > 0:
                    mean_val = values.mean()
                    if mean_val < limits['min'] or mean_val > limits['max']:
                        stats['out_of_spec'].append({
                            'parameter': col,
                            'mean': round(mean_val, 2),
                            'min_limit': limits['min'],
                            'max_limit': limits['max'],
                            'unit': limits['unit'],
                            'status': 'LOW' if mean_val < limits['min'] else 'HIGH'
                        })
                    # Trend analysis
                    if len(values) >= 3:
                        x = np.arange(len(values))
                        slope = np.polyfit(x, values, 1)[0]
                        stats['trends'][str(col)] = {
                            'slope': round(slope, 4),
                            'direction': 'Increasing' if slope > 0 else 'Decreasing',
                            'change_per_sample': round(slope, 4)
                        }
                break

    return stats


def generate_insights(stats_list):
    """Generate overall insights from all sheets."""
    insights = []
    alerts = []

    for stats in stats_list:
        if not stats:
            continue
        mat = stats['material']

        # Out of spec alerts
        for oos in stats['out_of_spec']:
            alerts.append({
                'type': 'danger',
                'message': f"📊 **{mat}** — **{oos['parameter']}** is {oos['status']} "
                          f"(Mean: {oos['mean']} {oos['unit']}, "
                          f"Limit: {oos['min_limit']}-{oos['max_limit']} {oos['unit']})"
            })

        # Trend insights
        for param, trend in stats['trends'].items():
            if abs(trend['slope']) > 0.5:
                direction = "📈 rising" if trend['slope'] > 0 else "📉 falling"
                alerts.append({
                    'type': 'warning',
                    'message': f"📊 **{mat}** — **{param}** is trending {direction} "
                              f"({trend['change_per_sample']} per sample)"
                })

        # General insights
        if stats['row_count'] > 1:
            insights.append(f"✅ **{mat}**: {stats['row_count']} samples analyzed")

    return insights, alerts


def create_trend_chart(df, material_type):
    """Create trend charts for numeric parameters."""
    numeric_df = df.select_dtypes(include=[np.number])
    if numeric_df.empty or len(numeric_df) < 2:
        return None

    # Use index as x-axis if no date column
    x_axis = df.index
    x_label = "Sample Number"

    # Check for date column
    for col in df.columns:
        if 'date' in str(col).lower():
            try:
                x_axis = pd.to_datetime(df[col], errors='coerce')
                x_label = "Date"
                break
            except:
                pass

    fig = make_subplots(
        rows=min(3, len(numeric_df.columns)),
        cols=1,
        subplot_titles=[str(c) for c in numeric_df.columns[:3]],
        vertical_spacing=0.1
    )

    for i, col in enumerate(numeric_df.columns[:3], 1):
        fig.add_trace(
            go.Scatter(x=x_axis, y=numeric_df[col], mode='lines+markers',
                      name=str(col), line=dict(width=2)),
            row=i, col=1
        )

        # Add specification lines if available
        specs = MATERIAL_KEYWORDS.get(material_type, {}).get('params', {})
        for param, limits in specs.items():
            if param.lower() in str(col).lower():
                fig.add_hline(y=limits['max'], line_dash="dash", line_color="red", 
                             annotation_text="Max", row=i, col=1)
                fig.add_hline(y=limits['min'], line_dash="dash", line_color="orange",
                             annotation_text="Min", row=i, col=1)
                break

    fig.update_layout(height=300*min(3, len(numeric_df.columns)), 
                     showlegend=False, template="plotly_white")
    return fig


def extract_data_from_image(image):
    """Extract text/data from uploaded image using OCR."""
    try:
        # Convert PIL to OpenCV format
        img_array = np.array(image)
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array

        # Preprocess for better OCR
        gray = cv2.resize(gray, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # OCR
        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(thresh, config=custom_config)

        # Try to extract tabular data
        data = pytesseract.image_to_data(thresh, config=custom_config, output_type=pytesseract.Output.DICT)

        return text, data
    except Exception as e:
        return f"OCR Error: {str(e)}", None


def parse_ocr_text(text):
    """Parse OCR text to extract structured data."""
    lines = text.split('\n')
    data_rows = []

    for line in lines:
        # Look for patterns like "Parameter: Value" or "Parameter  Value  Unit"
        parts = re.split(r'[:;|\t\s{2,}]', line.strip())
        parts = [p.strip() for p in parts if p.strip()]

        if len(parts) >= 2:
            # Try to identify parameter and value
            param = parts[0]
            for p in parts[1:]:
                try:
                    val = float(p.replace(',', '.'))
                    data_rows.append({'Parameter': param, 'Value': val})
                    break
                except:
                    continue

    if data_rows:
        return pd.DataFrame(data_rows)
    return None


def generate_word_report(all_results, insights, alerts):
    """Generate a Word document report."""
    doc = Document()

    # Title
    title = doc.add_heading('Cement Quality Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    if alerts:
        doc.add_paragraph('The following alerts require attention:')
        for alert in alerts[:10]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(alert['message'].replace('**', ''))
    else:
        doc.add_paragraph('All parameters are within normal operating ranges.')

    # Detailed Results
    doc.add_heading('Detailed Analysis', level=1)

    for result in all_results:
        if not result:
            continue
        mat = result['material']
        doc.add_heading(f'{mat} Analysis', level=2)
        doc.add_paragraph(f"Samples analyzed: {result['row_count']}")

        if result['out_of_spec']:
            doc.add_paragraph('Out of Specification Parameters:')
            for oos in result['out_of_spec']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{oos['parameter']}: {oos['mean']} {oos['unit']} ({oos['status']})")

        if result['trends']:
            doc.add_paragraph('Trends:')
            for param, trend in result['trends'].items():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{param}: {trend['direction']} ({trend['change_per_sample']} per sample)")

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def copy_button(text, key_suffix=""):
    """Create a copy-to-clipboard button using HTML/JS."""
    escaped_text = text.replace("`", "\`").replace("$", "\$")
    html = f"""
    <div style="margin: 10px 0;">
        <button id="copyBtn{key_suffix}" onclick="navigator.clipboard.writeText(`{escaped_text}`).then(() => {{
            const btn = document.getElementById('copyBtn{key_suffix}');
            btn.innerHTML = '✅ Copied!';
            btn.style.backgroundColor = '#28a745';
            setTimeout(() => {{
                btn.innerHTML = '📋 Copy Results to Clipboard';
                btn.style.backgroundColor = '#1f4e79';
            }}, 2000);
        }})" 
        style="background-color: #1f4e79; color: white; border: none; padding: 10px 20px; 
               border-radius: 8px; cursor: pointer; font-weight: 600; font-size: 14px;
               transition: all 0.3s;">
            📋 Copy Results to Clipboard
        </button>
    </div>
    """
    st.components.v1.html(html, height=60)


# =============================================================================
# SESSION STATE INITIALIZATION
# =============================================================================
if 'all_results' not in st.session_state:
    st.session_state.all_results = []
if 'insights' not in st.session_state:
    st.session_state.insights = []
if 'alerts' not in st.session_state:
    st.session_state.alerts = []
if 'extracted_df' not in st.session_state:
    st.session_state.extracted_df = None

# =============================================================================
# MAIN UI
# =============================================================================

st.markdown('<div class="main-header">🏭 Cement Clinker Quality Analyzer</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Multi-Material | Multi-Sheet | OCR-Enabled | Offline Ready</div>', unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.image("https://raw.githubusercontent.com/streamlit/streamlit/develop/frontend/assets/logo.png", width=100)
    st.title("⚙️ Controls")

    st.markdown("---")
    st.subheader("📋 Material Specifications")

    for mat, info in MATERIAL_KEYWORDS.items():
        with st.expander(f"📊 {mat} Specs"):
            for param, limits in info['params'].items():
                st.markdown(f"**{param}**: {limits['min']} - {limits['max']} {limits['unit']}")

    st.markdown("---")
    st.info("💡 **Tip**: Upload Excel with multiple sheets or use camera/image upload for OCR-based data extraction.")

    if st.session_state.alerts:
        st.markdown("---")
        st.subheader("🚨 Active Alerts")
        for alert in st.session_state.alerts[:5]:
            if alert['type'] == 'danger':
                st.error(alert['message'])
            elif alert['type'] == 'warning':
                st.warning(alert['message'])

# Main Tabs
tab1, tab2, tab3, tab4 = st.tabs([
    "📁 Excel Analysis", 
    "📷 Image & OCR", 
    "📊 Results & Trends", 
    "📋 Report & Export"
])

# =============================================================================
# TAB 1: EXCEL ANALYSIS
# =============================================================================
with tab1:
    st.header("📁 Multi-Sheet Excel Analysis")
    st.markdown("Upload an Excel file containing **Kiln Feed**, **Clinker**, and/or **Cement** data across multiple sheets.")

    col1, col2 = st.columns([2, 1])

    with col1:
        uploaded_excel = st.file_uploader(
            "Drop your Excel file here",
            type=['xlsx', 'xls'],
            help="Supports multi-sheet files with mixed material data"
        )

    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        analyze_btn = st.button("🔍 Analyze All Sheets", type="primary", use_container_width=True)
        clear_btn = st.button("🗑️ Clear Results", use_container_width=True)

        if clear_btn:
            st.session_state.all_results = []
            st.session_state.insights = []
            st.session_state.alerts = []
            st.rerun()

    if uploaded_excel and analyze_btn:
        with st.spinner("🔍 Scanning all sheets and detecting materials..."):
            try:
                xl_file = pd.ExcelFile(uploaded_excel)
                sheet_names = xl_file.sheet_names

                st.success(f"📑 Found **{len(sheet_names)}** sheet(s): {', '.join(sheet_names)}")

                all_results = []
                all_dfs = {}

                progress_bar = st.progress(0)

                for idx, sheet in enumerate(sheet_names):
                    progress_bar.progress((idx + 1) / len(sheet_names))

                    try:
                        df = pd.read_excel(xl_file, sheet_name=sheet)
                        df = clean_dataframe(df)

                        if df.empty:
                            st.warning(f"⚠️ Sheet '{sheet}' is empty after cleaning.")
                            continue

                        material = detect_material_type(df, sheet)

                        st.markdown(f"---")
                        st.subheader(f"📄 Sheet: '{sheet}' → Detected: **{material}**")

                        # Show preview
                        with st.expander(f"👁️ Preview Data ({len(df)} rows)", expanded=False):
                            st.dataframe(df, use_container_width=True, height=250)

                        # Analyze
                        stats = analyze_sheet(df, material)
                        if stats:
                            all_results.append(stats)
                            all_dfs[sheet] = {'df': df, 'material': material, 'stats': stats}

                            # Quick metrics
                            mcol1, mcol2, mcol3, mcol4 = st.columns(4)
                            numeric_df = df.select_dtypes(include=[np.number])

                            with mcol1:
                                st.metric("Samples", len(df))
                            with mcol2:
                                st.metric("Parameters", len(numeric_df.columns))
                            with mcol3:
                                oos_count = len(stats['out_of_spec'])
                                st.metric("⚠️ Out of Spec", oos_count, 
                                         delta="Critical" if oos_count > 0 else "OK",
                                         delta_color="inverse")
                            with mcol4:
                                trend_count = len(stats['trends'])
                                st.metric("📈 Trends", trend_count)

                    except Exception as e:
                        st.error(f"❌ Error processing sheet '{sheet}': {str(e)}")

                progress_bar.empty()

                # Generate insights
                insights, alerts = generate_insights(all_results)

                st.session_state.all_results = all_results
                st.session_state.insights = insights
                st.session_state.alerts = alerts
                st.session_state.all_dfs = all_dfs

                # Summary
                st.markdown("---")
                st.subheader("📊 Scan Summary")

                summary_col1, summary_col2, summary_col3 = st.columns(3)
                with summary_col1:
                    st.markdown('<div class="success-box">', unsafe_allow_html=True)
                    st.markdown(f"**✅ Sheets Processed:** {len([r for r in all_results if r])}")
                    st.markdown('</div>', unsafe_allow_html=True)
                with summary_col2:
                    st.markdown('<div class="alert-box">', unsafe_allow_html=True)
                    st.markdown(f"**⚠️ Alerts:** {len(alerts)}")
                    st.markdown('</div>', unsafe_allow_html=True)
                with summary_col3:
                    st.markdown('<div class="insight-box">', unsafe_allow_html=True)
                    st.markdown(f"**💡 Insights:** {len(insights)}")
                    st.markdown('</div>', unsafe_allow_html=True)

                if alerts:
                    st.subheader("🚨 Immediate Attention Required")
                    for alert in alerts:
                        if alert['type'] == 'danger':
                            st.markdown(f'<div class="danger-box">{alert["message"]}</div>', unsafe_allow_html=True)
                        else:
                            st.markdown(f'<div class="alert-box">{alert["message"]}</div>', unsafe_allow_html=True)

                if insights:
                    st.subheader("💡 Key Insights")
                    for insight in insights:
                        st.markdown(f'<div class="insight-box">{insight}</div>', unsafe_allow_html=True)

            except Exception as e:
                st.error(f"❌ Failed to process Excel file: {str(e)}")

    elif not uploaded_excel and analyze_btn:
        st.warning("⚠️ Please upload an Excel file first.")

# =============================================================================
# TAB 2: IMAGE & OCR
# =============================================================================
with tab2:
    st.header("📷 Image Upload & OCR Extraction")
    st.markdown("Upload lab report images, screenshots, or use your camera to extract data automatically.")

    img_col1, img_col2 = st.columns(2)

    with img_col1:
        st.subheader("📁 Upload Image")
        uploaded_image = st.file_uploader(
            "Upload lab report image",
            type=['png', 'jpg', 'jpeg', 'bmp', 'tiff'],
            help="Upload screenshots or photos of lab reports"
        )

    with img_col2:
        st.subheader("📸 Camera / Screenshot")
        camera_image = st.camera_input("Take a photo or screenshot")

    image_to_process = None
    source = ""

    if uploaded_image:
        image_to_process = Image.open(uploaded_image)
        source = "Uploaded Image"
    elif camera_image:
        image_to_process = Image.open(camera_image)
        source = "Camera/Screenshot"

    if image_to_process:
        st.markdown("---")
        st.subheader(f"🖼️ Processing: {source}")

        img_col, preview_col = st.columns([1, 2])

        with img_col:
            st.image(image_to_process, caption="Original Image", use_container_width=True)

        with preview_col:
            with st.spinner("🔍 Running OCR..."):
                text, raw_data = extract_data_from_image(image_to_process)

            st.subheader("📝 Extracted Text")
            st.text_area("Raw OCR Output", text, height=200)

            # Try to parse structured data
            parsed_df = parse_ocr_text(text)

            if parsed_df is not None and not parsed_df.empty:
                st.success(f"✅ Parsed {len(parsed_df)} data points from image!")
                st.dataframe(parsed_df, use_container_width=True)
                st.session_state.extracted_df = parsed_df

                # Allow adding to main results
                if st.button("➕ Add to Analysis Results", type="primary"):
                    st.session_state.all_results.append({
                        'material': 'OCR Extracted',
                        'row_count': len(parsed_df),
                        'columns': list(parsed_df.columns),
                        'numeric_summary': parsed_df.describe().to_dict(),
                        'out_of_spec': [],
                        'trends': {}
                    })
                    st.success("✅ Added to results! Check the Results tab.")
            else:
                st.info("ℹ️ Could not auto-parse tabular data. Please review the raw text above.")
                st.session_state.extracted_df = None

# =============================================================================
# TAB 3: RESULTS & TRENDS
# =============================================================================
with tab3:
    st.header("📊 Detailed Results & Trend Analysis")

    if not st.session_state.all_results:
        st.info("ℹ️ No data analyzed yet. Go to **Excel Analysis** or **Image & OCR** tab first.")
    else:
        # Material filter
        materials = list(set([r['material'] for r in st.session_state.all_results if r]))
        selected_mat = st.selectbox("🔍 Filter by Material", ["All"] + materials)

        for result in st.session_state.all_results:
            if not result:
                continue
            if selected_mat != "All" and result['material'] != selected_mat:
                continue

            mat = result['material']

            with st.expander(f"📊 {mat} — {result['row_count']} Samples", expanded=True):
                # Statistics table
                st.subheader("📈 Statistical Summary")
                if result['numeric_summary']:
                    summary_df = pd.DataFrame(result['numeric_summary']).T
                    st.dataframe(summary_df.style.format("{:.2f}"), use_container_width=True)

                # Out of spec
                if result['out_of_spec']:
                    st.subheader("⚠️ Out of Specification")
                    oos_df = pd.DataFrame(result['out_of_spec'])
                    st.dataframe(oos_df, use_container_width=True)

                # Trends
                if result['trends']:
                    st.subheader("📈 Trend Analysis")
                    trend_data = []
                    for param, trend in result['trends'].items():
                        trend_data.append({
                            'Parameter': param,
                            'Direction': trend['direction'],
                            'Change/Sample': trend['change_per_sample'],
                            'Status': '🔴 Significant' if abs(trend['change_per_sample']) > 0.5 else '🟢 Stable'
                        })
                    trend_df = pd.DataFrame(trend_data)
                    st.dataframe(trend_df, use_container_width=True)

                # Charts if dataframe available
                if 'all_dfs' in st.session_state:
                    for sheet, info in st.session_state.all_dfs.items():
                        if info['material'] == mat:
                            fig = create_trend_chart(info['df'], mat)
                            if fig:
                                st.plotly_chart(fig, use_container_width=True)
                            break

        # Copy results section
        st.markdown("---")
        st.subheader("📋 Copy Results")

        # Format results as text
        result_text = "CEMENT QUALITY ANALYSIS RESULTS\n"
        result_text += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        result_text += "=" * 50 + "\n\n"

        for result in st.session_state.all_results:
            if not result:
                continue
            result_text += f"MATERIAL: {result['material']}\n"
            result_text += f"Samples: {result['row_count']}\n"
            if result['out_of_spec']:
                result_text += "OUT OF SPEC:\n"
                for oos in result['out_of_spec']:
                    result_text += f"  - {oos['parameter']}: {oos['mean']} {oos['unit']} ({oos['status']})\n"
            if result['trends']:
                result_text += "TRENDS:\n"
                for param, trend in result['trends'].items():
                    result_text += f"  - {param}: {trend['direction']}\n"
            result_text += "\n" + "-" * 40 + "\n\n"

        st.code(result_text, language="text")
        copy_button(result_text, key_suffix="results")

# =============================================================================
# TAB 4: REPORT & EXPORT
# =============================================================================
with tab4:
    st.header("📋 Generate & Export Report")

    if not st.session_state.all_results:
        st.info("ℹ️ No data to report. Analyze data in previous tabs first.")
    else:
        st.subheader("📄 Report Preview")

        # Build preview
        report_md = f"""# Cement Quality Analysis Report
**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  
**Total Materials Analyzed:** {len([r for r in st.session_state.all_results if r])}

---

## Executive Summary
"""

        if st.session_state.alerts:
            report_md += "\n**⚠️ Alerts Requiring Attention:**\n\n"
            for alert in st.session_state.alerts[:10]:
                report_md += f"- {alert['message']}\n"
        else:
            report_md += "\n✅ All parameters are within normal operating ranges.\n"

        report_md += "\n---\n\n## Detailed Results\n"

        for result in st.session_state.all_results:
            if not result:
                continue
            report_md += f"\n### {result['material']}\n"
            report_md += f"- **Samples:** {result['row_count']}\n"

            if result['out_of_spec']:
                report_md += "- **⚠️ Out of Specification:**\n"
                for oos in result['out_of_spec']:
                    report_md += f"  - {oos['parameter']}: {oos['mean']} {oos['unit']} ({oos['status']})\n"

            if result['trends']:
                report_md += "- **📈 Trends:**\n"
                for param, trend in result['trends'].items():
                    report_md += f"  - {param}: {trend['direction']} ({trend['change_per_sample']} per sample)\n"

        st.markdown(report_md)

        # Copy report
        st.markdown("---")
        st.subheader("📋 Copy Full Report")
        st.code(report_md, language="markdown")
        copy_button(report_md, key_suffix="report")

        # Export buttons
        st.markdown("---")
        st.subheader("💾 Download Report")

        dcol1, dcol2 = st.columns(2)

        with dcol1:
            # Word report
            try:
                word_buffer = generate_word_report(
                    st.session_state.all_results,
                    st.session_state.insights,
                    st.session_state.alerts
                )
                st.download_button(
                    label="📄 Download Word Report (.docx)",
                    data=word_buffer,
                    file_name=f"Cement_Quality_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Word export error: {e}")

        with dcol2:
            # CSV combined export
            if 'all_dfs' in st.session_state:
                csv_buffer = io.StringIO()
                for sheet, info in st.session_state.all_dfs.items():
                    csv_buffer.write(f"\n# SHEET: {sheet} | MATERIAL: {info['material']}\n")
                    info['df'].to_csv(csv_buffer, index=False)

                st.download_button(
                    label="📊 Download Combined CSV",
                    data=csv_buffer.getvalue(),
                    file_name=f"Cement_Quality_Data_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )

# =============================================================================
# FOOTER
# =============================================================================
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #888; padding: 1rem;">
    <small>🏭 Cement Clinker Quality Analyzer | Offline Mode | Built with Streamlit</small>
</div>
""", unsafe_allow_html=True)
